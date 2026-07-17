import sys
import subprocess
import os

def install(package):
    pass
    # subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install("python-docx")
    import docx

try:
    from fpdf import FPDF
except ImportError:
    install("fpdf2")
    from fpdf import FPDF

import csv
from pathlib import Path

# Add backend to path so we can import from app
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from app.services.ingestion.parser_service import parser_service

def generate_csv(filepath):
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Equipment", "Status", "Date"])
        writer.writerow(["P-101", "Failed", "2026-05-10"])
        writer.writerow(["V-205", "Operational", "2026-05-11"])
    print(f"Generated {filepath}")

def generate_docx(filepath):
    doc = docx.Document()
    doc.add_heading('Maintenance Log for P-101', 0)
    doc.add_paragraph('On 2026-05-12, the pump seal was found to be leaking heavily.')
    doc.add_paragraph('Action taken: replaced seal and tested.')
    doc.save(filepath)
    print(f"Generated {filepath}")

def generate_pdf(filepath):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(40, 10, "Inspection Report")
    pdf.ln(10)
    pdf.set_font("helvetica", "", 12)
    pdf.cell(40, 10, "Equipment V-205 was inspected on 2026-05-13. No leaks found.")
    pdf.output(filepath)
    print(f"Generated {filepath}")

def test_parser(filepath):
    print(f"--- Testing {filepath} ---")
    try:
        md_text = parser_service.parse_document(Path(filepath))
        print("Success! Parsed Markdown:")
        print(md_text[:200]) # Print first 200 chars
        print("...")
        return True
    except Exception as e:
        print(f"Failed to parse {filepath}: {e}")
        return False

if __name__ == "__main__":
    csv_file = "test_doc.csv"
    docx_file = "test_doc.docx"
    pdf_file = "test_doc.pdf"
    
    generate_csv(csv_file)
    generate_docx(docx_file)
    generate_pdf(pdf_file)
    
    success_csv = test_parser(csv_file)
    success_docx = test_parser(docx_file)
    success_pdf = test_parser(pdf_file)
    
    print("\n--- Summary ---")
    print(f"CSV: {'PASS' if success_csv else 'FAIL'}")
    print(f"DOCX: {'PASS' if success_docx else 'FAIL'}")
    print(f"PDF: {'PASS' if success_pdf else 'FAIL'}")
