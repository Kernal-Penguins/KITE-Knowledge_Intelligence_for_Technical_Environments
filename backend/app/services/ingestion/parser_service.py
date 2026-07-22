"""
services/ingestion/parser_service.py
────────────────────────────────────
Docling wrapper for document parsing with lightweight fallbacks.
"""
from pathlib import Path
from app.infrastructure.logger import log


class ParserService:
    def __init__(self):
        self._converter = None
        self._docling_available = True

    @property
    def converter(self):
        # Lazy load because initialization can be heavy
        if self._converter is None and self._docling_available:
            try:
                log.info("parser_service.initializing_docling")
                from docling.document_converter import DocumentConverter
                self._converter = DocumentConverter()
                log.info("parser_service.initialized_docling")
            except Exception as e:
                log.warning("parser_service.docling_unavailable", error=str(e))
                self._docling_available = False
        return self._converter

    def parse_document(self, file_path: Path) -> str:
        """Parse a document into markdown representation using Docling with fallbacks."""
        log.info("parser_service.converting", file_path=str(file_path))
        file_path = Path(file_path)
        
        # Try Docling first if available
        if self.converter is not None:
            try:
                result = self.converter.convert(file_path)
                md_text = result.document.export_to_markdown()
                log.info("parser_service.converted_via_docling", file_path=str(file_path), md_length=len(md_text))
                return md_text
            except Exception as e:
                log.warning("parser_service.docling_failed_using_fallback", file_path=str(file_path), error=str(e))

        # Fallback parsers
        ext = file_path.suffix.lower()
        try:
            if ext == ".pdf":
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(str(file_path))
                text_pages = []
                for i, page in enumerate(pdf):
                    text_page = page.get_textpage()
                    text = text_page.get_text_range()
                    if text.strip():
                        text_pages.append(f"## Page {i+1}\n{text}")
                md_text = "\n\n".join(text_pages)
            elif ext in (".docx", ".doc"):
                import docx
                doc = docx.Document(str(file_path))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                md_text = "\n\n".join(paragraphs)
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    md_text = f.read()

            log.info("parser_service.converted_via_fallback", file_path=str(file_path), md_length=len(md_text))
            return md_text
        except Exception as fallback_err:
            log.error("parser_service.all_parsers_failed", file_path=str(file_path), error=str(fallback_err))
            raise fallback_err


parser_service = ParserService()

